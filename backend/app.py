from flask import Flask, request, jsonify, send_from_directory
from flask_cors import CORS
import sqlite3
import time
import os
from selenium import webdriver
import undetected_chromedriver as uc
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.chrome.options import Options
import google.generativeai as genai
from docx import Document  # For creating .doc files
from datetime import datetime  # To capture the current date and time

app = Flask(__name__)
CORS(app, resources={r"/*": {"origins": "*"}})

# Get the absolute path to the database file
db_path = os.path.abspath("data.db")
print(f"Database path: {db_path}")

# Directory to store .doc files
MINUTES_DIR = "minutes_files"
os.makedirs(MINUTES_DIR, exist_ok=True)  # Create the directory if it doesn't exist

# Configure Gemini API
genai.configure(api_key="AIzaSyDdphi9GL-bUV-4vtYKzw60GojymilnrYw")  # Replace with your Gemini API key
model = genai.GenerativeModel("gemini-1.5-pro-latest")  # Use the latest Gemini model

# Database Initialization
def init_db():
    try:
        conn = sqlite3.connect(db_path)
        cursor = conn.cursor()
        cursor.execute("""
            CREATE TABLE IF NOT EXISTS transcripts (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                meeting_id TEXT,
                meeting_name TEXT,
                meet_url TEXT,
                transcript TEXT,
                date TEXT,
                participants TEXT
            )
        """)
        cursor.execute("""
            CREATE TABLE IF NOT EXISTS minutes (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                meeting_id TEXT,
                meeting_name TEXT,
                meet_url TEXT,
                file_link TEXT,  -- Store the file path
                file_type TEXT   -- Store the file type (doc)
            )
        """)
        conn.commit()
        conn.close()
        print("Database initialized and 'transcripts' and 'minutes' tables created.")
    except Exception as e:
        print(f"Error initializing database: {e}")

init_db()  # Ensure DB is ready when the app starts

# Function to summarize the transcript using Gemini API
def summarize_transcript_with_gemini(transcript, date, participants):
    """Convert transcript into meeting minutes using Gemini API."""
    # Define the prompt
    prompt = f"""
    Format the following meeting transcript into structured minutes with these sections:

    Date: {date}
    Participants: {participants}

    Agenda:
    - [Summarize the main topic of the meeting]

    Discussion Points:
    1. [Key discussion points]
    2. [Additional discussion points]

    Decisions:
    - [Any decisions made during the meeting]

    Action Items:
    - [List of tasks assigned]

    Conclusion:
    - [Summary of the meeting]

    Transcript:
    {transcript}
    """

    try:
        # Call the Gemini API
        response = model.generate_content(prompt)
        return response.text
    except Exception as e:
        print(f"Error calling Gemini API: {e}")
        return None

# Function to generate minutes as a .doc file
def generate_minutes_doc(minutes, meeting_name, date, participants):
    """Generate minutes as a .doc file."""
    doc = Document()
    doc.add_heading(f"Meeting Minutes: {meeting_name}", level=1)
    doc.add_paragraph(f"Date: {date}")
    doc.add_paragraph(f"Participants: {participants}")
    doc.add_paragraph("")  # Add a blank line

    # Add the minutes content
    for line in minutes.split("\n"):
        doc.add_paragraph(line.strip())

    # Replace colons in the date to make the filename valid
    safe_date = date.replace(":", "-")  # Replace colons with hyphens
    doc_filename = f"minutes_{meeting_name}_{safe_date}.doc"
    doc_filepath = os.path.join(MINUTES_DIR, doc_filename)
    doc.save(doc_filepath)
    print(f"DOC file saved to {doc_filepath}")
    return doc_filepath


def start_meeting(meeting_id, meeting_name, url, participants):
    """Start a Google Meet session and extract captions."""
    # Automatically capture the current date and time
    current_date_time = datetime.now().strftime("%Y-%m-%d %H:%M:%S")

    options = Options()
    options.add_argument("--use-fake-ui-for-media-stream")  # Auto-approve mic/camera permissions
    options.add_argument("--disable-blink-features=AutomationControlled")  # Reduce bot detection
    options.add_argument("--disable-gpu")  # Fix GPU rendering error
    options.add_argument("--disable-software-rasterizer")
    options.add_argument("--no-sandbox")
    options.add_argument("--disable-dev-shm-usage")  # Prevent crashes in low-memory environments
    options.add_argument("--ignore-certificate-errors")
    options.add_argument("--remote-debugging-port=9222")  # Enable remote debugging
    options.add_argument("--disable-extensions")  # Disable extensions
    options.add_argument("--disable-popup-blocking")  # Disable popups
    options.add_argument("--disable-notifications")  # Disable notifications

    # Start undetected ChromeDriver
    driver = uc.Chrome(options=options, use_subprocess=True)
    driver.implicitly_wait(30)  # Increase implicit wait time

    print(f"Opening Google Meet: {url}")
    driver.get(url)
    print("Google Meet page loaded.")

    # Check if bot detection is bypassed
    is_bot = driver.execute_script("return navigator.webdriver")
    print(f"Bot Detection Bypassed: {is_bot}")

    if is_bot:
        print("Bot detection triggered. Try updating undetected-chromedriver.")
        driver.quit()
        return

    transcript = ""
    last_captions_text = ""  # Track the last captured text

    try:
        # Wait for captions container to load
        WebDriverWait(driver, 60).until(
            EC.presence_of_element_located((By.CSS_SELECTOR, 'div[aria-label="Captions"] div[jsname="tgaKEf"]'))
        )

        # Poll for captions
        while True:
            try:
                if not driver.session_id:
                    print("Browser session terminated.")
                    break

                captions = driver.find_elements(By.CSS_SELECTOR, 'div[aria-label="Captions"] div[jsname="tgaKEf"]')
                if captions:
                    current_captions_text = " ".join([c.text for c in captions])

                    # Split the captions text into lines
                    current_lines = current_captions_text.split(". ")
                    last_lines = last_captions_text.split(". ")

                    # Append only new lines to the transcript
                    for line in current_lines:
                        if line not in last_lines:
                            transcript += line + ". "
                            print("Captured Line:", line)

                    last_captions_text = current_captions_text  # Update the last captured text
                else:
                    print("No captions found.")

                time.sleep(5)  # Poll every 5 seconds
            except Exception as e:
                print("Error extracting captions:", e)
                break  # Exit loop if error occurs

    except Exception as e:
        print(f"Meet loading error: {e}")
        driver.save_screenshot("error.png")  # Take a screenshot on error

    # Summarize the transcript into meeting minutes using Gemini API
    minutes = summarize_transcript_with_gemini(transcript, current_date_time, participants)

    if not minutes:
        print("Failed to generate minutes. Using raw transcript as fallback.")
        minutes = transcript

    # Generate minutes as a .doc file
    doc_filepath = generate_minutes_doc(minutes, meeting_name, current_date_time, participants)

    # Store the transcript and minutes file link in the database
    store_transcript(meeting_id, meeting_name, url, transcript, current_date_time, participants, doc_filepath)

    # Keep the browser open for 30 seconds before quitting
    time.sleep(30)
    driver.quit()

def store_transcript(meeting_id, meeting_name, meet_url, transcript, date, participants, file_link):
    """Save extracted transcript and generated minutes in SQLite."""
    try:
        conn = sqlite3.connect(db_path)
        cursor = conn.cursor()
        
        # Store the transcript
        cursor.execute("""
            INSERT INTO transcripts (meeting_id, meeting_name, meet_url, transcript, date, participants) 
            VALUES (?, ?, ?, ?, ?, ?)
        """, (meeting_id, meeting_name, meet_url, transcript, date, participants))
        
        # Store the minutes file link in the database
        cursor.execute("""
            INSERT INTO minutes (meeting_id, meeting_name, meet_url, file_link, file_type) 
            VALUES (?, ?, ?, ?, ?)
        """, (meeting_id, meeting_name, meet_url, file_link, "doc"))
        
        conn.commit()
        conn.close()
        print("Transcript and minutes file link stored in the database.")
    except Exception as e:
        print(f"Error storing transcript and minutes: {e}")

@app.route('/start-meet', methods=['POST'])
def handle_meet():
    """API endpoint to receive Meet link and start transcript extraction."""
    data = request.json
    print("Received data:", data)  # Debugging: Print the request data

    if not data:
        return jsonify({"error": "No data provided"}), 400

    meeting_id = data.get("meeting_id")
    meeting_name = data.get("meeting_name")
    meet_url = data.get("url")
    participants = data.get("participants")

    if not all([meeting_id, meeting_name, meet_url, participants]):
        return jsonify({"error": "Missing required fields"}), 400

    start_meeting(meeting_id, meeting_name, meet_url, participants)
    return jsonify({"status": "Meeting started"})

@app.route('/download-minutes/<meeting_id>', methods=['GET'])
def download_minutes(meeting_id):
    """Download minutes file (.doc) for a given meeting ID."""
    try:
        conn = sqlite3.connect(db_path)
        cursor = conn.cursor()
        cursor.execute("""
            SELECT file_link FROM minutes WHERE meeting_id = ?
        """, (meeting_id,))
        result = cursor.fetchone()
        conn.close()

        if result:
            file_link = result[0]
            if os.path.exists(file_link):
                # Serve the file from the local directory
                return send_from_directory(
                    directory=os.path.dirname(file_link),
                    path=os.path.basename(file_link),
                    as_attachment=True
                )
            else:
                return jsonify({"error": "File not found"}), 404
        return jsonify({"error": "Minutes not found"}), 404
    except Exception as e:
        return jsonify({"error": str(e)}), 500

if __name__ == "__main__":
    app.run(debug=True)